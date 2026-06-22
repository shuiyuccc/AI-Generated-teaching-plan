#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
教案模板填充工具 - 保持原有格式，只填充单元格内容
"""

from docx import Document


def fill_cell_content(cell, text):
    """填充单元格内容，保持原有格式"""
    # 获取单元格中的所有段落
    paragraphs = cell.paragraphs
    
    if paragraphs:
        # 保留第一个段落，清除原有内容
        first_para = paragraphs[0]
        # 保留段落的格式，只替换文本
        for run in first_para.runs:
            run.text = ""
        # 如果有 runs，在第一个 run 中设置文本
        if first_para.runs:
            first_para.runs[0].text = text
        else:
            first_para.add_run(text)
        
        # 清除其他段落（如果有）
        for para in paragraphs[1:]:
            para.clear()


def fill_jiaoan():
    """填充教案模板，保持原有格式"""
    
    # 要填充的数据
    fill_data = {
        # 表格1 - 基本信息（左侧标签，右侧空白）
        "院    系": "信息工程学院",
        "授课班级": "2023级计算机1班",
        "专业名称": "计算机科学与技术",
        "课程名称": "Python程序设计基础",
        "授课教师": "张明华",
        
        # 表格2 - 详细信息
        "课题名称": "第四章 函数的定义与调用",
        "授课班级_2": "2023级计算机1班",  # 表格2中的授课班级
        "授课地点": "教学楼A301机房",
        "授课时间": "2024年3月15日 第3-4节",
        "授课学时": "2学时",
        "授课类型": "理论+实践",
        
        # 教学内容及学情分析
        "教学内容及学情分析": """本次课程主要讲解Python函数的基本概念、定义方法、参数传递以及返回值的使用。

具体包括：
1. 函数的定义与调用语法
2. 形参与实参的概念
3. 位置参数、关键字参数、默认参数
4. 可变参数（*args, **kwargs）
5. 函数的返回值
6. 变量的作用域

学情分析：
学生已掌握Python基础语法（变量、数据类型、流程控制），具备一定的编程思维能力。但对于模块化编程思想理解不深，需要通过实例引导学生理解函数封装的意义。班级学生编程基础参差不齐，需设计分层教学任务。""",
        
        # 教学目标
        "教学目标": """知识目标：理解函数的概念和作用；掌握函数的定义和调用方法；理解参数传递的几种方式；掌握return语句的使用。

能力目标：能够根据需求独立编写函数；能够正确使用不同类型的参数；能够分析并解决函数调用中的常见问题；培养模块化编程思维。

素质目标：培养代码复用意识；增强逻辑思维能力；培养严谨的程序设计习惯；提升团队协作解决问题的能力。""",
        
        # 教学重点
        "教学重点": "函数的定义与调用语法；参数传递机制；return返回值的使用。",
        
        # 教学难点
        "教学难点": "可变参数的理解与使用；变量作用域（局部变量与全局变量）；参数传递的内存机制。",
        
        # 教学方法与教学资源
        "教学方法与教学资源": """教学方法：
1. 案例驱动教学法：通过实际编程案例引入函数概念
2. 任务驱动教学法：布置分层编程任务，学生动手实践
3. 小组协作学习：分组讨论、互相调试代码
4. 演示法：教师演示关键代码和调试技巧

教学资源：
1. 多媒体课件（PPT）
2. Python 3.9+ 编程环境
3. 在线编程平台（Jupyter Notebook）
4. 课程案例代码库
5. 教学视频资源
6. 课后练习题库""",
        
        # 思政元素
        "思政元素": """1. 工匠精神：通过代码规范要求，培养学生精益求精的编程态度
2. 团队协作：函数模块化思想类比社会分工协作，培养学生的团队意识
3. 创新思维：鼓励学生优化算法，培养创新精神
4. 责任意识：强调代码质量的重要性，培养学生的责任心""",
        
        # 教学实施过程
        "教学环节": [
            {
                "环节": "课程导入",
                "时间": "10min",
                "教学内容": "回顾已学知识，通过'重复代码问题'引入函数概念。展示一段包含重复代码的程序，引导学生思考如何优化。",
                "教师活动": "提出问题，展示案例代码，引导学生思考",
                "学生活动": "观察代码，思考优化方案，回答问题"
            },
            {
                "环节": "新课讲授",
                "时间": "35min",
                "教学内容": "讲解函数定义语法、参数类型、返回值。通过实例演示函数的定义与调用过程。",
                "教师活动": "讲解概念，演示代码，板书重点",
                "学生活动": "听讲记录，跟随演示操作"
            },
            {
                "环节": "课堂练习",
                "时间": "25min",
                "教学内容": "学生完成3个函数编程练习：计算阶乘、判断素数、求最大值。",
                "教师活动": "布置任务，巡视指导，解答疑问",
                "学生活动": "独立编程，调试代码，举手提问"
            },
            {
                "环节": "作品展示",
                "时间": "10min",
                "教学内容": "选取2-3名学生展示代码，师生共同点评。",
                "教师活动": "组织学生展示，点评总结",
                "学生活动": "展示作品，相互学习"
            },
            {
                "环节": "课堂小结",
                "时间": "8min",
                "教学内容": "总结函数要点，强调易错点，布置课后作业。",
                "教师活动": "梳理知识，强调重点",
                "学生活动": "回顾整理，记录作业"
            },
            {
                "环节": "课后拓展",
                "时间": "2min",
                "教学内容": "预告下节课内容：递归函数",
                "教师活动": "布置预习任务",
                "学生活动": "记录预习内容"
            }
        ],
        
        # 课外作业
        "课外作业": """1. 完成教材P128页习题4-1、4-3、4-5
2. 编写一个计算器程序，使用函数实现加减乘除功能
3. 预习递归函数概念，思考递归与循环的区别
4. 在线平台完成本周编程练习（截止下周三）""",
        
        # 教学反思
        "教学反思": """（课后填写）

1. 教学目标达成情况：
2. 学生参与度分析：
3. 教学方法效果评估：
4. 存在问题及改进措施：
5. 下次教学调整方向："""
    }
    
    # 打开模板文档
    doc = Document("moban.docx")
    
    # 遍历所有表格
    table_idx = 0
    for table in doc.tables:
        table_idx += 1
        
        if table_idx == 1:
            # 表格1：院系统、授课班级、专业名称、课程名称、授课教师
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    label = cells[0].text.strip()
                    if "院" in label and "系" in label:
                        fill_cell_content(cells[1], fill_data["院    系"])
                    elif "授课班级" in label:
                        fill_cell_content(cells[1], fill_data["授课班级"])
                    elif "专业名称" in label:
                        fill_cell_content(cells[1], fill_data["专业名称"])
                    elif "课程名称" in label:
                        fill_cell_content(cells[1], fill_data["课程名称"])
                    elif "授课教师" in label:
                        fill_cell_content(cells[1], fill_data["授课教师"])
        
        elif table_idx == 2:
            # 表格2：详细信息
            row_idx = 0
            for row in table.rows:
                row_idx += 1
                cells = row.cells
                
                if row_idx == 1:
                    # 课题名称
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["课题名称"])
                
                elif row_idx == 2:
                    # 授课班级、授课地点
                    for i, cell in enumerate(cells):
                        text = cell.text.strip()
                        if "授课班级" in text:
                            if i + 1 < len(cells):
                                fill_cell_content(cells[i + 1], fill_data["授课班级_2"])
                        elif "授课地点" in text:
                            if i + 1 < len(cells):
                                fill_cell_content(cells[i + 1], fill_data["授课地点"])
                
                elif row_idx == 3:
                    # 授课时间、授课学时、授课类型
                    for i, cell in enumerate(cells):
                        text = cell.text.strip()
                        if "授课时间" in text:
                            if i + 1 < len(cells):
                                fill_cell_content(cells[i + 1], fill_data["授课时间"])
                        elif "授课学时" in text:
                            if i + 1 < len(cells):
                                fill_cell_content(cells[i + 1], fill_data["授课学时"])
                        elif "授课类型" in text or "此处填写理论" in text:
                            fill_cell_content(cell, fill_data["授课类型"])
                
                elif row_idx == 4:
                    # 教学内容及学情分析
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["教学内容及学情分析"])
                
                elif row_idx == 5:
                    # 教学目标
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["教学目标"])
                
                elif row_idx == 6:
                    # 教学重点
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["教学重点"])
                
                elif row_idx == 7:
                    # 教学难点
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["教学难点"])
                
                elif row_idx == 8:
                    # 教学方法与教学资源
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["教学方法与教学资源"])
                
                elif row_idx == 9:
                    # 思政元素
                    if len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["思政元素"])
        
        elif table_idx == 3:
            # 表格3：教学实施过程
            row_idx = 0
            环节_idx = 0
            for row in table.rows:
                row_idx += 1
                cells = row.cells
                
                # 跳过表头行（第1、2行）
                if row_idx <= 2:
                    continue
                
                # 填充教学环节
                if 环节_idx < len(fill_data["教学环节"]):
                    环节 = fill_data["教学环节"][环节_idx]
                    if len(cells) >= 4:
                        # 第一列：环节+时间
                        fill_cell_content(cells[0], f"{环节['环节']}{环节['时间']}")
                        # 第二列：教学内容
                        fill_cell_content(cells[1], 环节["教学内容"])
                        # 第三列：教师活动
                        fill_cell_content(cells[2], 环节["教师活动"])
                        # 第四列：学生活动
                        fill_cell_content(cells[3], 环节["学生活动"])
                    环节_idx += 1
                
                # 检查是否是课外作业行
                if len(cells) >= 1:
                    first_cell_text = cells[0].text.strip()
                    if "课外作业" in first_cell_text and len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["课外作业"])
                    elif "教学反思" in first_cell_text and len(cells) >= 2:
                        fill_cell_content(cells[1], fill_data["教学反思"])
    
    # 保存文档
    output_path = "moban_filled.docx"
    doc.save(output_path)
    print(f"教案填充完成！已保存至: {output_path}")
    print("所有内容已填充，原有格式保持不变。")


if __name__ == "__main__":
    fill_jiaoan()
