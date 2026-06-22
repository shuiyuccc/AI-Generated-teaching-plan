#!/usr/bin/env python3
"""
分析教案模板结构 - 包含合并单元格信息
生成供大模型分析的JSON文件
"""
from docx import Document
import json


def get_cell_merge_info(cell):
    """
    获取单元格的合并信息
    """
    tc = cell._tc
    tcPr = tc.tcPr
    
    # 检查水平合并
    gridSpan = tcPr.gridSpan if tcPr is not None else None
    
    # 检查垂直合并
    vMerge = tcPr.vMerge if tcPr is not None else None
    
    merge_info = {
        "是否合并单元格": False,
        "合并类型": None,
        "水平合并": {
            "是否合并": False,
            "跨度": 1
        },
        "垂直合并": {
            "是否合并": False,
            "类型": None,
            "跨度": 1
        }
    }
    
    # 水平合并分析
    if gridSpan is not None:
        val = gridSpan.val
        if val is not None and int(val) > 1:
            merge_info["是否合并单元格"] = True
            merge_info["合并类型"] = "horizontal"
            merge_info["水平合并"]["是否合并"] = True
            merge_info["水平合并"]["跨度"] = int(val)
    
    # 垂直合并分析
    if vMerge is not None:
        vMerge_val = vMerge.val
        if vMerge_val is None:
            # 这是垂直合并的延续单元格
            merge_info["是否合并单元格"] = True
            merge_info["垂直合并"]["是否合并"] = True
            merge_info["垂直合并"]["类型"] = "continue"
            if merge_info["合并类型"] == "horizontal":
                merge_info["合并类型"] = "both"
            else:
                merge_info["合并类型"] = "vertical"
        elif vMerge_val == "restart":
            # 垂直合并的起始单元格
            merge_info["是否合并单元格"] = True
            merge_info["垂直合并"]["是否合并"] = True
            merge_info["垂直合并"]["类型"] = "restart"
            if merge_info["合并类型"] == "horizontal":
                merge_info["合并类型"] = "both"
            else:
                merge_info["合并类型"] = "vertical"
    
    return merge_info


def calculate_merge_matrix(table):
    """
    计算整个表格的合并矩阵
    返回每个单元格的合并信息以及它属于哪个合并区域
    """
    rows = list(table.rows)
    row_count = len(rows)
    
    # 创建合并信息矩阵
    merge_matrix = []
    for row_idx, row in enumerate(rows):
        row_info = []
        for col_idx, cell in enumerate(row.cells):
            merge_info = get_cell_merge_info(cell)
            merge_info["行索引"] = row_idx
            merge_info["列索引"] = col_idx
            row_info.append(merge_info)
        merge_matrix.append(row_info)
    
    # 创建合并区域映射表
    # key: (row, col), value: 合并区域信息
    merge_regions = {}
    
    for row_idx in range(row_count):
        col_idx = 0
        while col_idx < len(merge_matrix[row_idx]):
            cell_info = merge_matrix[row_idx][col_idx]
            
            # 检查是否是水平合并的起始单元格
            if cell_info["水平合并"]["是否合并"] and cell_info["水平合并"]["跨度"] > 1:
                h_span = cell_info["水平合并"]["跨度"]
                
                # 记录这个合并区域
                region_id = f"区域_行{row_idx}_列{col_idx}"
                merge_regions[(row_idx, col_idx)] = {
                    "区域ID": region_id,
                    "类型": "水平合并",
                    "起始行": row_idx,
                    "起始列": col_idx,
                    "结束行": row_idx,
                    "结束列": col_idx + h_span - 1,
                    "行跨度": 1,
                    "列跨度": h_span,
                    "包含单元格": [(row_idx, c) for c in range(col_idx, col_idx + h_span)]
                }
                
                # 为合并区域内的所有单元格标记所属区域
                for c in range(col_idx, col_idx + h_span):
                    merge_matrix[row_idx][c]["所属合并区域"] = region_id
                    merge_matrix[row_idx][c]["合并区域信息"] = merge_regions[(row_idx, col_idx)]
                    
                    # 标记是否是左上角（填充位置）
                    if c == col_idx:
                        merge_matrix[row_idx][c]["是左上角"] = True
                        merge_matrix[row_idx][c]["填充位置"] = {"行": row_idx, "列": col_idx}
                    else:
                        merge_matrix[row_idx][c]["是左上角"] = False
                        merge_matrix[row_idx][c]["填充位置"] = {"行": row_idx, "列": col_idx}
                
                col_idx += h_span
            else:
                # 普通单元格或垂直合并（你的模板没有垂直合并）
                if "所属合并区域" not in merge_matrix[row_idx][col_idx]:
                    merge_matrix[row_idx][col_idx]["所属合并区域"] = None
                    merge_matrix[row_idx][col_idx]["是左上角"] = True
                    merge_matrix[row_idx][col_idx]["填充位置"] = {"行": row_idx, "列": col_idx}
                col_idx += 1
    
    return merge_matrix, merge_regions


def analyze_cell(cell, row_idx, col_idx, merge_info):
    """分析单个单元格的完整信息"""
    cell_info = {
        "位置": {"行": row_idx, "列": col_idx},
        "文本内容": cell.text,
        "段落数量": len(cell.paragraphs),
        "合并信息": {
            "是否合并单元格": merge_info["是否合并单元格"],
            "合并类型": merge_info["合并类型"],
            "水平合并": merge_info["水平合并"],
            "垂直合并": merge_info["垂直合并"]
        },
        "所属合并区域": merge_info.get("所属合并区域"),
        "是左上角": merge_info.get("是左上角", True),
        "填充位置": merge_info.get("填充位置", {"行": row_idx, "列": col_idx}),
        "代码访问方式": f"table.cell({row_idx}, {col_idx})"
    }
    
    # 添加填充建议
    if merge_info.get("是左上角"):
        if merge_info["是否合并单元格"]:
            region = merge_info.get("合并区域信息", {})
            span_info = f"(合并了{region.get('列跨度', 1)}列)" if region else ""
            cell_info["填充建议"] = f"✅ 在此处填充数据 {span_info}"
        else:
            cell_info["填充建议"] = "✅ 在此处填充数据"
    else:
        fill_pos = merge_info.get("填充位置", {})
        cell_info["填充建议"] = f"➡️ 这是合并区域的一部分，应该填充到 table.cell({fill_pos['行']}, {fill_pos['列']})"
    
    # 如果有合并区域信息，添加详细的合并范围
    if merge_info.get("合并区域信息"):
        region = merge_info["合并区域信息"]
        cell_info["合并范围"] = {
            "说明": f"这个单元格与行{region['起始行']}列{region['起始列']}到行{region['结束行']}列{region['结束列']}合并",
            "合并区域包含的单元格": [f"行{r}-列{c}" for r, c in region["包含单元格"]]
        }
    
    return cell_info


def analyze_table(table, table_idx):
    """分析表格的完整结构"""
    rows = list(table.rows)
    row_count = len(rows)
    col_count = len(table.columns)
    
    # 计算合并矩阵
    merge_matrix, merge_regions = calculate_merge_matrix(table)
    
    table_info = {
        "表格索引": table_idx,
        "表格名称": get_table_name(table_idx),
        "尺寸": {"行数": row_count, "列数": col_count},
        "合并单元格统计": {
            "合并区域数量": len(merge_regions),
            "合并区域列表": []
        },
        "行详情": []
    }
    
    # 统计合并区域
    for (start_row, start_col), region in merge_regions.items():
        table_info["合并单元格统计"]["合并区域列表"].append({
            "区域ID": region["区域ID"],
            "位置": f"行{start_row}-列{start_col}",
            "类型": region["类型"],
            "范围": f"从 行{region['起始行']}列{region['起始列']} 到 行{region['结束行']}列{region['结束列']}",
            "列跨度": region["列跨度"],
            "填充位置": f"table.cell({start_row}, {start_col})",
            "包含单元格": [f"行{r}-列{c}" for r, c in region["包含单元格"]]
        })
    
    # 分析每一行
    for row_idx, row in enumerate(rows):
        row_info = {
            "行索引": row_idx,
            "单元格": []
        }
        
        for col_idx, cell in enumerate(row.cells):
            merge_info = merge_matrix[row_idx][col_idx]
            cell_info = analyze_cell(cell, row_idx, col_idx, merge_info)
            row_info["单元格"].append(cell_info)
        
        table_info["行详情"].append(row_info)
    
    return table_info


def get_table_name(table_idx):
    """获取表格名称"""
    names = {0: "封面信息表", 1: "教案内容表", 2: "教学实施过程表"}
    return names.get(table_idx, f"表格{table_idx}")


def generate_filling_guide(doc_info):
    """生成填充指南"""
    guide = {
        "说明": "此指南告诉大模型应该填充到哪些单元格（只有左上角才需要填充）",
        "重要提示": "只需要填充标记为'是左上角': true 的单元格",
        "表格填充指南": []
    }
    
    for table in doc_info["表格详情"]:
        table_guide = {
            "表格索引": table["表格索引"],
            "表格名称": table["表格名称"],
            "需要填充的单元格": []
        }
        
        for row in table["行详情"]:
            for cell in row["单元格"]:
                # 只记录左上角且为空的单元格
                if cell.get("是左上角") and cell["文本内容"].strip() == "":
                    fill_instruction = {
                        "位置": f"行{cell['位置']['行']}-列{cell['位置']['列']}",
                        "代码": cell["代码访问方式"]
                    }
                    
                    if cell["合并信息"]["是否合并单元格"]:
                        region_info = cell.get("合并范围", {})
                        fill_instruction["合并信息"] = {
                            "类型": cell["合并信息"]["合并类型"],
                            "水平跨度": cell["合并信息"]["水平合并"]["跨度"],
                            "垂直跨度": cell["合并信息"]["垂直合并"]["跨度"],
                            "合并范围说明": region_info.get("说明", "")
                        }
                    
                    table_guide["需要填充的单元格"].append(fill_instruction)
        
        guide["表格填充指南"].append(table_guide)
    
    return guide


def analyze_document(template_path: str):
    """分析整个文档结构"""
    doc = Document(template_path)
    
    doc_info = {
        "文档路径": template_path,
        "段落数量": len(doc.paragraphs),
        "表格数量": len(doc.tables),
        "使用说明": {
            "摘要": "此JSON文件描述了教案模板的结构，包含合并单元格信息",
            "关键概念": {
                "合并单元格": "多个单元格合并成一个大单元格",
                "左上角": "合并单元格的第一个单元格，是唯一需要填充的位置",
                "合并区域": "被合并的一组单元格，共享同一个填充位置"
            },
            "如何使用": {
                "步骤1": "查找标记为 '是左上角': true 的单元格",
                "步骤2": "这些单元格就是需要填充的位置",
                "步骤3": "使用 '代码访问方式' 中的代码进行填充",
                "步骤4": "如果是合并单元格，填充一次会自动应用到整个合并区域"
            }
        },
        "表格详情": []
    }
    
    # 分析每个表格
    for table_idx, table in enumerate(doc.tables):
        table_info = analyze_table(table, table_idx)
        doc_info["表格详情"].append(table_info)
    
    # 生成填充指南
    doc_info["填充指南"] = generate_filling_guide(doc_info)
    
    return doc_info


def main():
    template_path = "moban.docx"
    
    print("[分析] 正在分析模板文件的合并单元格...")
    doc_info = analyze_document(template_path)
    
    # 保存完整分析到文件
    output_file = "template_structure_with_merge.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(doc_info, f, ensure_ascii=False, indent=2)
    
    print(f"[完成] 完整分析已保存到: {output_file}")
    print(f"[信息] 文档包含 {doc_info['表格数量']} 个表格")
    
    # 打印简要的合并信息
    print("\n[合并单元格摘要]")
    for table in doc_info["表格详情"]:
        stats = table["合并单元格统计"]
        print(f"\n表格 {table['表格索引']} ({table['表格名称']}):")
        print(f"  合并区域数量: {stats['合并区域数量']}")
        if stats['合并区域列表']:
            print("  合并区域详情:")
            for region in stats['合并区域列表']:
                print(f"    - {region['位置']}: {region['类型']}, 列跨度={region['列跨度']}")
                print(f"      填充位置: {region['填充位置']}")


if __name__ == "__main__":
    main()
