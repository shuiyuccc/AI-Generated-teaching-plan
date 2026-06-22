"""
教案模板管理器
支持用户上传自定义 Word 模板，自动分析结构并动态填充
"""
import os
import json
import logging
from docx import Document

logger = logging.getLogger(__name__)

# 存储已上传的模板
_uploaded_templates = {}  # {session_id: {path, structure, ...}}


def analyze_template(template_path):
    """
    分析 Word 模板结构，找出所有标签-值对的映射

    Returns:
        {
            'basic_info': [{'label': '院系', 'value_cell': (table, row, col)}, ...],
            'detail_info': [{'label': '课题名称', ...}, ...],
            'sections': ['教学内容及学情分析', '教学目标', ...],
            'teaching_steps_table': {'table_idx': N, 'header_rows': N, 'columns': [...]},
            'all_labels': ['院系', '授课班级', ...]
        }
    """
    doc = Document(template_path)
    structure = {
        'basic_info': [],
        'detail_info': [],
        'sections': [],
        'teaching_steps_table': None,
        'all_labels': [],
        'table_count': len(doc.tables)
    }

    # 常见的教案标签关键词
    section_keywords = [
        '教学内容及学情分析', '学情分析', '教学内容',
        '教学目标', '教学重点', '教学难点',
        '教学方法', '教学资源', '教学方法与教学资源',
        '思政元素', '课程思政',
        '课外作业', '课后作业', '教学反思',
        '教学实施', '教学过程', '教学环节'
    ]

    info_keywords = [
        '院系', '授课班级', '专业名称', '课程名称', '授课教师',
        '课题名称', '授课地点', '授课时间', '授课学时', '授课类型'
    ]

    for table_idx, table in enumerate(doc.tables):
        # 分析每个表格
        first_row_text = ''
        if table.rows:
            first_cells = table.rows[0].cells
            first_row_text = ' '.join([c.text.strip() for c in first_cells])

        # 检查是否是教学环节表格（有多列标题如"教学环节/时间/教学内容/教师活动/学生活动"）
        is_teaching_table = False
        teaching_columns = []
        for row in table.rows[:3]:  # 检查前3行
            row_texts = [c.text.strip() for c in row.cells]
            joined = ' '.join(row_texts)
            if any(kw in joined for kw in ['教学环节', '教师活动', '学生活动', '教学实施', '教学过程']):
                is_teaching_table = True
                teaching_columns = row_texts
                break

        if is_teaching_table:
            structure['teaching_steps_table'] = {
                'table_idx': table_idx,
                'columns': teaching_columns,
                'header_rows': 2  # 默认跳过前2行表头
            }
            continue

        # 解析标签-值对
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                text = cell.text.strip()
                if not text:
                    continue

                # 检查是否是信息标签
                for kw in info_keywords:
                    if kw in text and len(text) < 20:  # 标签一般比较短
                        # 找到相邻的值单元格
                        value_cell = None
                        if col_idx + 1 < len(cells):
                            value_cell = (table_idx, row_idx, col_idx + 1)
                        elif col_idx + 2 < len(cells):
                            value_cell = (table_idx, row_idx, col_idx + 2)

                        if value_cell:
                            entry = {
                                'label': text,
                                'keyword': kw,
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'col_idx': col_idx,
                                'value_cell': value_cell
                            }
                            # 区分基本信息表还是详细信息表
                            if table_idx == 0:
                                structure['basic_info'].append(entry)
                            else:
                                structure['detail_info'].append(entry)
                            structure['all_labels'].append(text)

                # 检查是否是段落型标签（如"教学内容及学情分析"）
                for kw in section_keywords:
                    if kw in text and len(cells) >= 2:
                        # 这类标签后面通常跟着一个大的内容单元格
                        value_col = 1 if col_idx == 0 else col_idx + 1
                        if value_col < len(cells):
                            entry = {
                                'label': text,
                                'keyword': kw,
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'col_idx': col_idx,
                                'value_cell': (table_idx, row_idx, value_col)
                            }
                            structure['sections'].append(entry)
                            structure['all_labels'].append(text)

    logger.info(f"模板分析完成: {len(structure['all_labels'])} 个标签, "
                f"{len(structure['sections'])} 个段落区, "
                f"教学环节表: {'有' if structure['teaching_steps_table'] else '无'}")

    return structure


def save_template_info(session_id, filepath, structure):
    """保存模板信息"""
    _uploaded_templates[session_id] = {
        'path': filepath,
        'structure': structure,
        'filename': os.path.basename(filepath)
    }


def get_template_info(session_id):
    """获取模板信息"""
    return _uploaded_templates.get(session_id)


def build_prompt_from_template(course_info, structure):
    """根据模板结构动态构建提示词"""
    labels = structure['all_labels']
    sections = [s['label'] for s in structure['sections']]
    has_teaching_table = bool(structure['teaching_steps_table'])

    topic = course_info.get('课题名称', '未指定')
    hours = course_info.get('授课学时', '')
    lesson_type = course_info.get('授课类型', '')

    # 参考文档
    ref_docs = course_info.get('参考文档', [])
    ref_text = ''
    if ref_docs:
        ref_text = '\n'.join([f"参考资料《{d.get('filename', '')}》：\n{d.get('content', '')[:2000]}" for d in ref_docs])

    sections_desc = '\n'.join([f"  - \"{s}\"" for s in sections])

    # 构建 JSON 示例
    teaching_section_json = ''
    if has_teaching_table:
        teaching_section_json = ',\n  "教学环节": [{"环节": "...", "时间": "Xmin", "教学内容": "...", "教师活动": "...", "学生活动": "..."}]'

    json_template = """{
  "sections": {
    "教学内容及学情分析": "详细分析...",
    "教学目标": "知识目标：...\\n能力目标：...\\n素质目标：...",
    "教学重点": "...",
    "教学难点": "...",
    "教学方法与教学资源": "...",
    "思政元素": "..."
  }""" + teaching_section_json + """,
  "课外作业": "...",
  "教学反思": "（课后填写）\\n1. 教学目标达成情况：\\n2. 学生参与度分析：\\n3. 教学方法效果评估：\\n4. 存在问题及改进措施："
}"""

    teaching_hint = ''
    if has_teaching_table:
        teaching_hint = '\n  - 教学环节（包含环节名称、时间、教学内容、教师活动、学生活动）'

    prompt = f"""你是一位资深教师，请根据以下课程信息生成一份完整教案。

【课程信息】
- 课题名称：{topic}
- 授课学时：{hours}
- 授课类型：{lesson_type}

{ref_text}

【需要生成的内容】
{sections_desc}{teaching_hint}
  - 课外作业
  - 教学反思模板

请按以下 JSON 格式输出（不要偷懒，每部分都要详细写300字以上）：

```json
{json_template}
```

直接输出 JSON，不要加其他说明。"""
    return prompt


def fill_template_dynamic(template_path, output_path, llm_data, course_info, structure):
    """
    根据分析出的模板结构动态填充 Word 文档
    """
    doc = Document(template_path)

    # 1. 填充基本信息标签
    info_mapping = {
        '院系': '院系', '授课班级': '授课班级', '专业名称': '专业名称',
        '课程名称': '课程名称', '授课教师': '授课教师',
        '课题名称': '课题名称', '授课地点': '授课地点',
        '授课时间': '授课时间', '授课学时': '授课学时', '授课类型': '授课类型'
    }

    for entry in structure['basic_info'] + structure['detail_info']:
        table_idx, row_idx, col_idx = entry['value_cell']
        table = doc.tables[table_idx]
        cell = table.rows[row_idx].cells[col_idx]

        # 尝试从 course_info 匹配
        value = None
        for key, field in info_mapping.items():
            if key in entry['label']:
                value = course_info.get(field, '')
                break

        if not value:
            value = ''

        _fill_cell_text(cell, str(value))

    # 2. 填充分段内容
    sections_data = llm_data.get('sections', {})
    for entry in structure['sections']:
        table_idx, row_idx, col_idx = entry['value_cell']
        table = doc.tables[table_idx]
        cell = table.rows[row_idx].cells[col_idx]

        # 匹配 LLM 返回的内容
        content = ''
        for key, val in sections_data.items():
            if key in entry['label'] or entry['label'] in key:
                content = val
                break
        if not content and entry['keyword'] in sections_data:
            content = sections_data[entry['keyword']]

        _fill_cell_text(cell, str(content) if content else '')

    # 3. 填充教学环节表格
    if structure['teaching_steps_table'] and '教学环节' in llm_data:
        steps_table_info = structure['teaching_steps_table']
        table = doc.tables[steps_table_info['table_idx']]
        steps = llm_data['教学环节']
        header_rows = steps_table_info.get('header_rows', 2)

        step_idx = 0
        for row_idx, row in enumerate(table.rows):
            if row_idx < header_rows:
                continue
            if step_idx < len(steps):
                step = steps[step_idx]
                cells = row.cells
                col_texts = [
                    f"{step.get('环节', '')}\n{step.get('时间', '')}",
                    step.get('教学内容', ''),
                    step.get('教师活动', ''),
                    step.get('学生活动', '')
                ]
                for i, text in enumerate(col_texts):
                    if i < len(cells):
                        _fill_cell_text(cells[i], text)
                step_idx += 1

            # 检查是否有课外作业和教学反思行
            first_text = row.cells[0].text.strip() if row.cells else ''
            if '课外作业' in first_text and len(row.cells) >= 2:
                _fill_cell_text(row.cells[1], llm_data.get('课外作业', ''))
            elif '教学反思' in first_text and len(row.cells) >= 2:
                _fill_cell_text(row.cells[1], llm_data.get('教学反思', ''))

    # 4. 填充没有明确标签的行（课外作业/教学反思可能不在教学环节表中）
    for table in doc.tables:
        for row in table.rows:
            first_text = row.cells[0].text.strip() if row.cells else ''
            if len(row.cells) >= 2:
                if '课外作业' in first_text and not row.cells[1].text.strip():
                    _fill_cell_text(row.cells[1], llm_data.get('课外作业', ''))
                elif '教学反思' in first_text and not row.cells[1].text.strip():
                    _fill_cell_text(row.cells[1], llm_data.get('教学反思', ''))

    doc.save(output_path)
    logger.info(f"动态填充完成: {output_path}")


def _fill_cell_text(cell, text):
    """填充单元格文本，保留原有格式"""
    paragraphs = cell.paragraphs
    if paragraphs:
        first_para = paragraphs[0]
        for run in first_para.runs:
            run.text = ""
        if first_para.runs:
            first_para.runs[0].text = str(text)
        else:
            first_para.add_run(str(text))
        for para in paragraphs[1:]:
            para.clear()
